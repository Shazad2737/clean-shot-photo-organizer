"""
Script to convert markdown documentation to a single DOCX file.
"""
import os
import re
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE


def create_docx_report():
    """Create a comprehensive DOCX report from markdown files."""
    
    # Create document
    doc = Document()
    
    # Set up styles
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    
    # Title page
    title = doc.add_heading('CLEAN SHOT', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_paragraph('AI Photo Organizer')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle.runs[0]
    subtitle_run.font.size = Pt(24)
    subtitle_run.font.color.rgb = RGBColor(124, 58, 237)  # Purple
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    project_info = doc.add_paragraph('Project Report Documentation')
    project_info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    project_info.runs[0].font.size = Pt(18)
    
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()
    
    # Project details table
    table = doc.add_table(rows=5, cols=2)
    table.style = 'Table Grid'
    
    details = [
        ('Project Name', 'CLEAN SHOT - AI Photo Organizer'),
        ('Version', '1.0.0'),
        ('License', 'MIT'),
        ('Platform', 'Windows, Linux, macOS'),
        ('Technology', 'Python, PySide6, OpenCV'),
    ]
    
    for i, (key, value) in enumerate(details):
        table.rows[i].cells[0].text = key
        table.rows[i].cells[1].text = value
    
    doc.add_page_break()
    
    # Define sections to include
    sections = [
        ('1_INTRODUCTION.md', '1. Introduction'),
        ('2_SYSTEM_OVERVIEW.md', '2. System Overview'),
        ('3_METHODOLOGY.md', '3. Methodology'),
        ('4_DESIGN.md', '4. Design'),
        ('5_TOOLS_AND_TECHNOLOGIES.md', '5. Tools & Technologies'),
        ('6_IMPLEMENTATION.md', '6. Implementation'),
        ('7_SAMPLE_CODE.md', '7. Sample Code'),
        ('8_SCREENSHOTS.md', '8. Screenshots'),
        ('9_TESTING.md', '9. Testing'),
        ('10_CONCLUSION.md', '10. Conclusion'),
    ]
    
    docs_dir = os.path.dirname(os.path.abspath(__file__))
    
    for filename, section_title in sections:
        filepath = os.path.join(docs_dir, filename)
        
        if os.path.exists(filepath):
            with open(filepath, 'r', encoding='utf-8') as f:
                content = f.read()
            
            # Process content
            process_markdown_to_docx(doc, content)
            doc.add_page_break()
        else:
            doc.add_heading(section_title, 1)
            doc.add_paragraph(f'[Content from {filename} not found]')
            doc.add_page_break()
    
    # Save document
    output_path = os.path.join(docs_dir, 'CLEAN_SHOT_Project_Report.docx')
    doc.save(output_path)
    print(f'✅ Document saved to: {output_path}')
    return output_path


def process_markdown_to_docx(doc, content):
    """Convert markdown content to DOCX elements."""
    
    lines = content.split('\n')
    in_code_block = False
    code_content = []
    in_table = False
    table_rows = []
    
    i = 0
    while i < len(lines):
        line = lines[i]
        
        # Code blocks
        if line.strip().startswith('```'):
            if in_code_block:
                # End code block
                if code_content:
                    code_text = '\n'.join(code_content)
                    p = doc.add_paragraph()
                    run = p.add_run(code_text)
                    run.font.name = 'Consolas'
                    run.font.size = Pt(9)
                    p.paragraph_format.left_indent = Inches(0.5)
                code_content = []
                in_code_block = False
            else:
                in_code_block = True
            i += 1
            continue
        
        if in_code_block:
            code_content.append(line)
            i += 1
            continue
        
        # Tables
        if '|' in line and not line.strip().startswith('```'):
            if not in_table:
                in_table = True
                table_rows = []
            
            # Skip separator lines
            if re.match(r'^[\|\s\-:]+$', line):
                i += 1
                continue
            
            cells = [cell.strip() for cell in line.split('|')[1:-1]]
            if cells:
                table_rows.append(cells)
            i += 1
            
            # Check if next line is not a table
            if i >= len(lines) or '|' not in lines[i]:
                # Create table
                if table_rows:
                    create_table(doc, table_rows)
                in_table = False
                table_rows = []
            continue
        
        # Headings
        if line.startswith('# '):
            doc.add_heading(clean_text(line[2:]), 1)
        elif line.startswith('## '):
            doc.add_heading(clean_text(line[3:]), 2)
        elif line.startswith('### '):
            doc.add_heading(clean_text(line[4:]), 3)
        elif line.startswith('#### '):
            doc.add_heading(clean_text(line[5:]), 4)
        # Bullet points
        elif line.strip().startswith('- ') or line.strip().startswith('* '):
            text = line.strip()[2:]
            p = doc.add_paragraph(clean_text(text), style='List Bullet')
        # Numbered lists
        elif re.match(r'^\d+\.\s', line.strip()):
            text = re.sub(r'^\d+\.\s', '', line.strip())
            p = doc.add_paragraph(clean_text(text), style='List Number')
        # Checkboxes
        elif line.strip().startswith('- [ ]') or line.strip().startswith('- [x]') or line.strip().startswith('- [/]'):
            text = line.strip()[6:]
            checked = '[x]' in line or '[/]' in line
            prefix = '☑ ' if checked else '☐ '
            p = doc.add_paragraph(prefix + clean_text(text), style='List Bullet')
        # Blockquotes
        elif line.strip().startswith('>'):
            text = line.strip()[1:].strip()
            p = doc.add_paragraph(clean_text(text))
            p.paragraph_format.left_indent = Inches(0.5)
            p.runs[0].font.italic = True
        # Regular paragraph
        elif line.strip():
            p = doc.add_paragraph(clean_text(line))
        
        i += 1


def create_table(doc, rows):
    """Create a table in the document."""
    if not rows:
        return
    
    num_cols = max(len(row) for row in rows)
    table = doc.add_table(rows=len(rows), cols=num_cols)
    table.style = 'Table Grid'
    
    for i, row in enumerate(rows):
        for j, cell in enumerate(row):
            if j < num_cols:
                table.rows[i].cells[j].text = clean_text(cell)
                # Bold header row
                if i == 0:
                    for paragraph in table.rows[i].cells[j].paragraphs:
                        for run in paragraph.runs:
                            run.font.bold = True
    
    doc.add_paragraph()  # Add space after table


def clean_text(text):
    """Clean markdown formatting from text."""
    # Remove bold/italic markers
    text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
    text = re.sub(r'\*(.+?)\*', r'\1', text)
    text = re.sub(r'__(.+?)__', r'\1', text)
    text = re.sub(r'_(.+?)_', r'\1', text)
    # Remove inline code
    text = re.sub(r'`(.+?)`', r'\1', text)
    # Remove links but keep text
    text = re.sub(r'\[([^\]]+)\]\([^\)]+\)', r'\1', text)
    # Remove images
    text = re.sub(r'!\[([^\]]*)\]\([^\)]+\)', r'[Image: \1]', text)
    return text.strip()


if __name__ == '__main__':
    create_docx_report()
